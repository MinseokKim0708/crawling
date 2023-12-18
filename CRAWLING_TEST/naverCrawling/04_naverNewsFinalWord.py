import requests
from bs4 import BeautifulSoup
import time
import pyautogui
from docx import Document

keyword = pyautogui.prompt("검색어를 입력하세요")
lastPage = int(pyautogui.prompt("몇 페이지까지 크롤링 할까요?"))

document = Document()

mainUrl = f'https://search.naver.com/search.naver?where=news&sm=tab_jum&query={keyword}'
pageNum = 1

for i in range(1, lastPage*10, 10):
    print(f"==============={pageNum}페이지 크롤링 중 입니다.====================")
    response = requests.get(mainUrl)
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')

    articles = soup.select("div.info_group")

    for article in articles:
        links = article.select("a.info")
        if len(links) >= 2:
            articleUrl = links[1].attrs['href']
            response = requests.get(articleUrl, headers={'User-agent':'Mozila/5.0'})
            html = response.text
            soup = BeautifulSoup(html, 'html.parser')

            if 'entertain' in response.url:
                title = soup.select_one(".end_tit")
                mainText = soup.select_one("#articeBody")
            
            elif 'sports' in response.url:
                title = soup.select_one("h4.title")
                mainText = soup.select_one("#newsEndContents")
                divs = mainText.select("div")
                for div in divs:
                    div.decompose()
                paragraphs = mainText.select("p")
                for p in paragraphs:
                    p.decompose()
            
            else:
                title = soup.select_one("#title_area")
                mainText = soup.select_one("#newsct_article")

            print("==========링크==========\n", articleUrl)
            print("==========기사제목==========\n", title.text.strip())
            print("==========본문내용==========\n", mainText.text.strip())

            document.add_heading(title.text.strip(), level=0)
            document.add_paragraph(articleUrl)
            document.add_paragraph(mainText.text.strip())
            time.sleep(0.3)

    isLastPage = soup.select_one("a.btn_next").attrs['aria-disabled']
    if isLastPage == 'true':
        print("마지막 페이지 입니다.")
        break
    pageNum += 1 

document.save(f"CRAWLING_TEST\\naverCrawling\\{keyword}_result.docx")